from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "复赛作品说明文档.docx"
GENERATED = ROOT / "images" / "generated"

# narrative_proposal preset + named CJK font override.
FONT_CN = "DengXian"
FONT_LATIN = "Aptos"
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(22, 36, 55)
MUTED = RGBColor(91, 103, 119)
PINK = RGBColor(229, 70, 119)
PURPLE = RGBColor(113, 73, 214)
LIGHT = "F4F6F9"
LIGHT_BLUE = "EEF4FB"
BORDER = "DCE3EC"
WHITE = RGBColor(255, 255, 255)

PIL_FONT = Path("/System/Library/Fonts/Hiragino Sans GB.ttc")
PIL_BOLD = Path("/System/Library/Fonts/Hiragino Sans GB.ttc")


SECTIONS = {
    "problem": (
        "高校学生与年轻职场人长期处在学业、实习、关系、家庭期待和自我评价叠加的压力场中。"
        "WHO 2025 年资料显示，全球近 1/7 人口正与心理障碍共处[1]；但 DayGlyph 聚焦的是更广泛、非临床的日常自我观察需求。"
        "传统长日记在疲惫时启动成本高，单选心情或“好/坏”量表又会压平欣慰与疲惫、期待与担忧并存的真实体验。"
        "细粒度情绪研究表明，人们区分相近情绪的能力具有重要观察价值，重复的日常记录也可能促进更细致的情绪辨认[4][5]。"
        "因此，用户真正需要的不是更多诊断、任务与打卡，而是一个足够轻、可信任、能把“一句话”沉淀为长期线索的入口。"
    ),
    "competitors": (
        "Daylio 以表情和活动标签实现两步记录，优势是轻量统计，但复杂中文感受仍需用户先自我归类[8]；"
        "Reflectly、Stoic 强调 AI 提问、模板与练习，容易把体验组织为“完成任务”；"
        "Apple 健康 State of Mind 具备系统级记录与关联图表，但入口仍是情绪滑杆、词项和影响因素选择[9]。"
        "DayGlyph 选择不同的技术路线：自然语言先进入 1～8 个细粒度情绪，再投影到 12 个锚点与 VAD 连续空间；"
        "同一向量同时驱动双视觉签名、微行动、回声和长期趋势。与“记录后给图表”相比，它把理解、表达、行动与回顾连接为闭环；"
        "日记、版本、行动和回声默认留在本机，用户可跳过、可删除、不受连续签到惩罚。"
    ),
    "feasibility": (
        "DayGlyph 已形成端到端 iOS 原型。Seed 2.0 Lite 将输入解析为情绪词、所属情绪族、强度、置信度、原文证据及 VAD 等连续维度；"
        "GenerationSchemaValidator 再校验 1～8 项数量、受控词库、数值区间、禁止诊断词与风险冲突。"
        "映射层按情绪族聚合并归一化权重：w_k = Σ intensity_i / Σ intensity_j，得到 12 维锚点分布；"
        "z=(V,A,D) 与 w 共同进入固定参数的 GlyphSignature，把效价映射到轨迹方向、唤醒度映射到节奏密度、掌控感映射到核心尺度，"
        "再以 hash(文本, 日期) 约束微扰，保证同文同日稳定。DayGenerationOrchestrator 以状态机编排“安全预检—结构化理解—Schema 校验—"
        "文字先展示—双图并行—本地保存—单图重试”，任一图片失败不阻塞另一项。SwiftData 保存日记与生成版本，图片落本地；"
        "趋势和回声由纯逻辑聚合器计算，并配套 Swift Testing、UI 测试、VoiceOver、减少动态与二维降级。"
    ),
    "innovation": (
        "DayGlyph 的创新不是“给日记加 AI”，而是构建一套可解释、可复现、可演进的情绪编译系统。"
        "其一，“离散锚点 + 连续 VAD”双空间同时保留情绪语义与可计算几何，避免单标签丢失混合情绪；细粒度文本情绪研究也支持多类别表达的必要性[6]。"
        "其二，提出“语义—视觉同构”：鸡尾酒与日星球共享色板、温度、密度和光线约束，形成可追溯的双视觉签名，而非两次无关生图。"
        "其三，以“记录—微行动—延迟回声—统计发现”构成反馈闭环；当前仅在同类样本不少于 3 次时呈现分布，不把相关性写成疗效。"
        "其四，规划受控自进化记忆：借鉴分层记忆与反思式管理研究[7]，将每日事件、跨期摘要和用户确认偏好分层，"
        "只对可验证节点执行 ADD、UPDATE、DECAY、DELETE；记忆可查看、可撤回，原文不作为隐藏人格标签。"
        "其五，把安全做成算法分支：高风险输入在生图前短路，普通低落不被医疗化，所有历史版本保持稳定。"
    ),
    "prospect": (
        "Grand View Research 预计全球心理健康 App 市场将由 2024 年约 74.8 亿美元增长至 2030 年约 175.2 亿美元[10]。"
        "在增长与信任赤字并存的市场中，DayGlyph 不与诊疗产品竞争，而切入高校学生、独立创作者和年轻职场人的日常情绪记录，"
        "并可作为咨询前自述、校园活动和个人年度回顾的低压力辅助工具。近期完善真机稳定性、离线降级与可解释评测；"
        "中期扩展 iPad、Apple Watch、年度报告及端侧模型；长期形成“中文情绪词库 + 可解释视觉编码 + 用户可控记忆协议”的技术资产。"
        "其商业价值来自高复访的个人数据沉淀与跨设备订阅，社会价值则是让更多人在需要专业帮助之前或之外，"
        "拥有一个不评判、可持续、守住数据主权的自我观察入口。"
    ),
}


def set_run_font(run, size=None, bold=None, color=None, italic=None, font_cn=FONT_CN):
    run.font.name = font_cn
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:eastAsia"), font_cn)
    r_pr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    r_pr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_pr.rFonts.set(qn("w:cs"), font_cn)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para(paragraph, *, before=0, after=8, line=1.333, align=None, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True
    if align is not None:
        paragraph.alignment = align


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_WIDTH_DXA
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, 8.2, color=MUTED)


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_CN
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_CN
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = DARK_BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT_CN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(header, after=0, line=1.0)
    set_run_font(header.add_run("DayGlyph 日迹  |  复赛作品说明文档"), 8.2, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(footer, after=0, line=1.0)
    set_run_font(footer.add_run("DayGlyph  ·  "), 8.2, color=MUTED)
    add_field(footer, "PAGE")


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=5, line=1.0)
    set_run_font(p.add_run("第十一届（2026年）中国高校计算机大赛—移动应用创新赛"), 9.5, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=7, line=1.0)
    set_run_font(p.add_run("复赛作品说明文档"), 12, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=5, line=1.08)
    set_run_font(p.add_run("DayGlyph 日迹"), 25, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=7, line=1.12)
    set_run_font(p.add_run("可解释情绪计算与受控自进化记忆驱动的长期自我观察系统"), 14, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=10, line=1.0)
    set_run_font(p.add_run("细粒度情绪建模  /  双视觉签名  /  反馈记忆  /  本地优先  /  非医疗化"), 9.2, bold=True, color=PINK)


def add_body(doc, text, *, after=8):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.333, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_run_font(p.add_run(text), 10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    set_run_font(p.add_run(text), 8.5, color=MUTED)


def set_picture_alt_text(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", description)
    doc_pr.set("descr", description)


def add_figure(doc, path, width_in, caption, *, before=2, after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    set_para(p, before=before, after=2, line=1.0)
    picture = p.add_run().add_picture(str(path), width=Inches(width_in))
    set_picture_alt_text(picture, caption)
    add_caption(doc, caption)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(after)


def add_two_image_figure(doc, left, right, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    set_para(p, before=2, after=2, line=1.0)
    left_picture = p.add_run().add_picture(str(left), width=Inches(2.45))
    set_picture_alt_text(left_picture, f"{caption}：左侧为情绪鸡尾酒")
    set_run_font(p.add_run("      "), 6)
    right_picture = p.add_run().add_picture(str(right), width=Inches(2.45))
    set_picture_alt_text(right_picture, f"{caption}：右侧为日星球")
    add_caption(doc, caption)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1600, 7760])
    set_table_borders(table, color="D8E3F1", size="5")
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=0, line=1.1)
    set_run_font(p.add_run(label), 9.2, bold=True, color=BLUE)
    p = table.cell(0, 1).paragraphs[0]
    set_para(p, after=0, line=1.15)
    set_run_font(p.add_run(text), 9.2, bold=True, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_matrix(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for j, text in enumerate(headers):
        set_cell_shading(table.cell(0, j), "E8EEF5")
        p = table.cell(0, j).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, after=0, line=1.05)
        set_run_font(p.add_run(text), font_size, bold=True, color=DARK_BLUE)
    for ridx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            if ridx % 2 == 0:
                set_cell_shading(cells[j], "FAFBFD")
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_para(p, after=0, line=1.1)
            set_run_font(p.add_run(text), font_size, bold=(j == 0), color=INK)
    set_table_geometry(table, widths)
    set_table_borders(table)
    p = doc.add_paragraph()
    set_para(p, before=4, after=6, line=1.0)
    return table


def make_competitor_strip():
    out = GENERATED / "semifinal-competitor-strip.png"
    canvas = Image.new("RGB", (1800, 650), "#F7F9FC")
    draw = ImageDraw.Draw(canvas)
    title_font = ImageFont.truetype(str(PIL_BOLD), 46)
    label_font = ImageFont.truetype(str(PIL_BOLD), 28)
    draw.text((70, 42), "同类产品界面样本", font=title_font, fill="#18314F")
    groups = [
        ("Daylio", ROOT / "竞品分析 image" / "dailo3.png"),
        ("Reflectly", ROOT / "竞品分析 image" / "reflectly1.png"),
        ("Stoic", ROOT / "竞品分析 image" / "stoic3.png"),
        ("DayGlyph", ROOT / "images" / "今日界面 4.png"),
    ]
    x_positions = [85, 510, 935, 1360]
    for (label, path), x in zip(groups, x_positions):
        card = Image.new("RGB", (330, 470), "white")
        image = Image.open(path).convert("RGB")
        image.thumbnail((290, 380))
        card.paste(image, ((330 - image.width) // 2, 58))
        canvas.paste(card, (x, 120))
        bbox = draw.textbbox((0, 0), label, font=label_font)
        draw.text((x + (330 - (bbox[2] - bbox[0])) / 2, 600), label, font=label_font, fill="#30587C")
    out.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(out, quality=95)
    return out


def rounded_card(draw, xy, fill, outline="#D8E2F0", radius=28, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered_text(draw, xy, text, font, fill, line_gap=8):
    x0, y0, x1, y1 = xy
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=font)
        widths.append(b[2] - b[0])
        heights.append(b[3] - b[1])
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y0 + (y1 - y0 - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x0 + (x1 - x0 - w) / 2, y), line, font=font, fill=fill)
        y += h + line_gap


def make_algorithm_diagram():
    out = GENERATED / "semifinal-algorithm-pipeline.png"
    canvas = Image.new("RGB", (2000, 900), "white")
    draw = ImageDraw.Draw(canvas)
    title = ImageFont.truetype(str(PIL_BOLD), 58)
    card_title = ImageFont.truetype(str(PIL_BOLD), 34)
    card_body = ImageFont.truetype(str(PIL_FONT), 25)
    formula = ImageFont.truetype(str(PIL_FONT), 25)
    draw.text((90, 54), "DayGlyph 可解释情绪编译链路", font=title, fill="#132B49")
    cards = [
        ("01  输入", "一句自然语言\n保留原文证据"),
        ("02  结构化", "1～8 种情绪\n强度 + 置信度"),
        ("03  校验", "词库 / 范围\n安全 / 禁止词"),
        ("04  投影", "12 锚点 + VAD\n权重归一化"),
        ("05  编译", "GlyphSignature\n稳定视觉参数"),
        ("06  闭环", "双图 / 微行动\n回声 / 趋势"),
    ]
    fills = ["#F6F8FC", "#EEF6FF", "#F1F8F5", "#F3EFFF", "#FFF3F5", "#FFF8EB"]
    x0 = 70
    y0 = 185
    w = 275
    h = 360
    gap = 48
    for idx, ((heading, body), fill) in enumerate(zip(cards, fills)):
        x = x0 + idx * (w + gap)
        rounded_card(draw, (x, y0, x + w, y0 + h), fill)
        draw_centered_text(draw, (x + 15, y0 + 40, x + w - 15, y0 + 118), heading, card_title, "#2E74B5")
        draw_centered_text(draw, (x + 18, y0 + 145, x + w - 18, y0 + 320), body, card_body, "#26384D")
        if idx < len(cards) - 1:
            arrow_x = x + w + 12
            draw.line((arrow_x, y0 + 180, arrow_x + 25, y0 + 180), fill="#7C69E8", width=8)
            draw.polygon(
                [(arrow_x + 25, y0 + 164), (arrow_x + 45, y0 + 180), (arrow_x + 25, y0 + 196)],
                fill="#7C69E8",
            )
    rounded_card(draw, (250, 625, 1750, 810), "#F8FAFD", outline="#DDE6F2")
    draw_centered_text(
        draw,
        (280, 642, 1720, 790),
        "w(k) = sum intensity(i) / sum intensity(j)     ·     G = F(z, w, hash(text, date))\n同文同日稳定；模型输出不能绕过 Schema、风险与版本边界",
        formula,
        "#234A72",
        line_gap=20,
    )
    canvas.save(out, quality=96)
    return out


def make_memory_diagram():
    out = GENERATED / "semifinal-memory-evolution.png"
    canvas = Image.new("RGB", (1900, 920), "#FBFCFF")
    draw = ImageDraw.Draw(canvas)
    title = ImageFont.truetype(str(PIL_BOLD), 56)
    hfont = ImageFont.truetype(str(PIL_BOLD), 34)
    bfont = ImageFont.truetype(str(PIL_FONT), 26)
    tagfont = ImageFont.truetype(str(PIL_BOLD), 23)
    draw.text((85, 50), "受控自进化记忆：让系统演进，但不把用户变成黑箱画像", font=title, fill="#182E4B")
    layers = [
        ("事件记忆｜已实现", "每日原文、情绪向量、视觉签名、行动与回声\n按 generationID 与版本保存；用户可删除", "#EDF5FF"),
        ("语义记忆｜已实现", "月 / 季 / 年趋势、关键词、情绪构成与回声分布\n纯逻辑聚合；样本不足不生成结论", "#F2EFFF"),
        ("偏好记忆｜复赛演进", "只写入用户确认或多次反馈支持的偏好节点\n可查看、可撤回、可衰减；不推断人格", "#FFF3F6"),
    ]
    y = 185
    for heading, body, fill in layers:
        rounded_card(draw, (115, y, 1380, y + 180), fill, outline="#D7E2F0")
        draw.text((155, y + 28), heading, font=hfont, fill="#2E74B5")
        draw.multiline_text((155, y + 86), body, font=bfont, fill="#2D4158", spacing=12)
        y += 220
    rounded_card(draw, (1450, 200, 1800, 760), "white", outline="#D7E2F0")
    draw_centered_text(draw, (1470, 225, 1780, 310), "记忆操作符", hfont, "#7049D6")
    ops = [("ADD", "新增证据"), ("UPDATE", "合并确认"), ("DECAY", "降低过期权重"), ("DELETE", "用户撤回")]
    oy = 335
    for op, desc in ops:
        rounded_card(draw, (1495, oy, 1755, oy + 78), "#F7F4FF", outline="#E5DCF8", radius=18, width=2)
        draw_centered_text(draw, (1505, oy + 5, 1745, oy + 73), f"{op}  {desc}", tagfont, "#5637B7")
        oy += 100
    draw.text((124, 844), "约束：本地优先 · 最小必要数据 · 结论需证据 · 相关不等于因果 · 全生命周期可审计", font=bfont, fill="#52677E")
    canvas.save(out, quality=96)
    return out


def add_cover(doc):
    add_title_block(doc)
    add_figure(
        doc,
        ROOT / "DayGlyph 宣传图.png",
        4.15,
        "图 1  DayGlyph 复赛作品视觉：一句话记录、情绪印记、微行动与延迟回声",
        before=0,
        after=0,
    )


def build():
    GENERATED.mkdir(parents=True, exist_ok=True)
    competitor_strip = make_competitor_strip()
    algorithm_diagram = make_algorithm_diagram()
    memory_diagram = make_memory_diagram()

    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    doc.add_page_break()

    add_heading(doc, "1、问题背景与用户分析（约200字）")
    add_figure(
        doc,
        ROOT / "images" / "generated" / "problem-background-ai.png",
        6.2,
        "图 2  用户问题结构：叠加压力、模糊感受与低门槛自我观察需求",
        before=0,
        after=5,
    )
    add_body(doc, SECTIONS["problem"])
    add_callout(
        doc,
        "需求定义",
        "不是再做一张情绪量表，而是在表达能力最低的时候，仍能完成一次低成本、可积累的自我观察。",
    )

    add_heading(doc, "2、相关竞品分析（约200字）")
    add_figure(doc, competitor_strip, 6.35, "图 3  同类产品界面样本：从手动选择、引导练习到自然语言情绪计算", before=0, after=5)
    add_matrix(
        doc,
        ["产品", "输入方式", "情绪表示", "长期闭环", "核心差异"],
        [
            ["Daylio", "表情 + 活动", "离散心情", "统计与相关图表", "极低记录成本"],
            ["Reflectly / Stoic", "引导问题 / 模板", "叙事或自评", "练习与回顾", "结构化自我照护"],
            ["Apple State of Mind", "滑杆 + 词项", "标准化状态", "历史与关联", "系统级健康入口"],
            ["DayGlyph", "一句自然语言", "1～8 情绪 → 12 锚点 + VAD", "双视觉 + 行动回声 + 宇宙", "可解释、本地优先、低压力"],
        ],
        [1350, 1800, 2100, 2050, 2060],
        font_size=8.1,
    )
    add_body(doc, SECTIONS["competitors"], after=3)
    add_figure(
        doc,
        ROOT / "images" / "readme" / "feature-today.jpg",
        6.2,
        "图 4  DayGlyph 今日页：情绪结构、双视觉、微行动与支持内容在同一链路呈现",
        before=2,
        after=2,
    )

    add_heading(doc, "3、可行性分析（约300字）")
    add_figure(doc, algorithm_diagram, 6.35, "图 5  从自然语言到长期回顾的可解释情绪编译算法", before=0, after=5)
    add_matrix(
        doc,
        ["阶段", "核心输入", "确定性约束", "可验证输出"],
        [
            ["结构化理解", "一句话文本", "1～8 项、受控词库、原文证据", "细粒度情绪 + 连续维度"],
            ["锚点投影", "情绪族与强度", "聚合、归一化、范围 clamp", "12 维权重 + VAD"],
            ["视觉编译", "z、w、稳定 seed", "固定公式与共享视觉方向", "鸡尾酒 + 日星球"],
            ["反馈聚合", "行动与回声", "n≥3、只报告分布、不推断因果", "个人历史发现"],
        ],
        [1500, 2200, 3150, 2510],
        font_size=8.3,
    )
    add_body(doc, SECTIONS["feasibility"], after=5)
    add_figure(
        doc,
        ROOT / "images" / "generated" / "feasibility-architecture-ai.png",
        6.25,
        "图 6  工程闭环：状态机编排、双图并发、本地持久化与可访问降级",
        before=0,
        after=3,
    )
    add_matrix(
        doc,
        ["工程能力", "代码级证据", "失败 / 降级策略"],
        [
            ["生成编排", "DayGenerationOrchestrator 状态机", "单图独立重试；重启后补齐未保存资源"],
            ["可信输出", "SchemaValidator + SafetyPrescreen", "字段非法回退；高风险在生图前短路"],
            ["历史稳定", "generationID + 版本 + 稳定 seed", "历史不静默重算；同文同日稳定"],
            ["可访问性", "VoiceOver + 减少动态 + 二维列表", "RealityKit 不可用时仍可完整浏览"],
        ],
        [1650, 3400, 4310],
        font_size=8.3,
    )

    doc.add_page_break()
    add_heading(doc, "4、App创新点（约300字）")
    add_two_image_figure(
        doc,
        ROOT / "DayGlyph" / "DemoAssets" / "demo-01-cocktail.jpeg",
        ROOT / "DayGlyph" / "DemoAssets" / "demo-01-planet.jpeg",
        "图 7  同一次情绪结构生成的鸡尾酒—日星球双视觉签名示例",
    )
    add_body(doc, SECTIONS["innovation"], after=5)
    add_figure(doc, memory_diagram, 6.35, "图 8  受控自进化记忆：已实现基础与复赛演进边界", before=0, after=5)

    doc.add_page_break()
    add_heading(doc, "5、应用前景（约200字）")
    add_figure(
        doc,
        ROOT / "images" / "generated" / "prospect-roadmap-ai.png",
        6.25,
        "图 9  产品路径：核心人群验证、跨设备延伸与长期技术资产",
        before=0,
        after=5,
    )
    add_matrix(
        doc,
        ["阶段", "产品目标", "算法 / 工程重点", "价值验证"],
        [
            ["近期", "iPhone 核心闭环稳定", "离线降级、可解释评测、真机性能", "完成率与回访意愿"],
            ["中期", "iPad / Watch / 年度报告", "端侧模型、跨设备同步边界", "留存与付费意愿"],
            ["长期", "用户可控的个人情绪资产", "中文词库、视觉编码、记忆协议", "可信个性化壁垒"],
        ],
        [1200, 2600, 3260, 2300],
        font_size=8.4,
    )
    add_body(doc, SECTIONS["prospect"], after=5)

    doc.add_page_break()
    add_heading(doc, "参考文献")
    refs = [
        (
            "World Health Organization. Mental disorders[EB/OL], 2025.",
            "WHO",
            "https://www.who.int/en/news-room/fact-sheets/detail/mental-disorders",
        ),
        (
            "Russell J A. A circumplex model of affect[J]. Journal of Personality and Social Psychology, 1980, 39(6): 1161-1178.",
            "DOI",
            "https://doi.org/10.1037/h0077714",
        ),
        (
            "Bradley M M, Lang P J. Measuring emotion: The Self-Assessment Manikin and the semantic differential[J]. Journal of Behavior Therapy and Experimental Psychiatry, 1994, 25(1): 49-59.",
            "PubMed",
            "https://pubmed.ncbi.nlm.nih.gov/7962581/",
        ),
        (
            "Hoemann K, Barrett L F, Quigley K S. Emotional Granularity Increases With Intensive Ambulatory Assessment[J]. Frontiers in Psychology, 2021.",
            "PMC",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC8355493/",
        ),
        (
            "Nook E C, et al. The Nonlinear Development of Emotion Differentiation[J]. Psychological Science, 2018, 29(8): 1346-1357.",
            "PMC",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC6088506/",
        ),
        (
            "Demszky D, et al. GoEmotions: A Dataset of Fine-Grained Emotions[C]. ACL, 2020: 4040-4054.",
            "ACL Anthology",
            "https://aclanthology.org/2020.acl-main.372/",
        ),
        (
            "Tan Z, et al. In Prospect and Retrospect: Reflective Memory Management for Long-term Personalized Dialogue Agents[C]. ACL, 2025: 8416-8439.",
            "ACL Anthology",
            "https://aclanthology.org/2025.acl-long.413/",
        ),
        ("Daylio. Mood Tracker and Micro Diary official website[EB/OL].", "Daylio", "https://daylio.net/"),
        (
            "Apple. Log your state of mind in Health on iPhone[EB/OL].",
            "Apple Support",
            "https://support.apple.com/guide/iphone/log-your-state-of-mind-iph6a6decb13/ios",
        ),
        (
            "Grand View Research. Mental Health Apps Market Size, Share & Trends Analysis Report[EB/OL], 2026.",
            "Grand View Research",
            "https://www.grandviewresearch.com/industry-analysis/mental-health-apps-market-report",
        ),
    ]
    for i, (text, label, url) in enumerate(refs, start=1):
        p = doc.add_paragraph()
        set_para(p, before=0, after=2, line=1.08)
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        set_run_font(p.add_run(f"[{i}] "), 8.2, bold=True, color=INK)
        set_run_font(p.add_run(text + " "), 8.2, color=INK)
        add_hyperlink(p, label, url)

    doc.core_properties.title = "DayGlyph 日迹复赛作品说明文档"
    doc.core_properties.subject = "中国高校计算机大赛—移动应用创新赛复赛作品说明"
    doc.core_properties.author = "DayGlyph 项目组"
    doc.core_properties.keywords = "DayGlyph, 情绪计算, 生成式视觉, 自进化记忆, iOS"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
