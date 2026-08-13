from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "初赛作品说明文档.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(30, 96, 145)
DARK = RGBColor(26, 38, 52)
MUTED = RGBColor(93, 103, 118)
BORDER = "E4E8F0"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para(paragraph, *, before=0, after=6, line=1.2, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def cell_margins(table, top=80, start=80, bottom=80, end=80):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1E6091")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(footer, after=0, line=1.0)
    run = footer.add_run("DayGlyph 日迹 · 初赛作品说明文档")
    set_run_font(run, 8.2, color=MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=6, line=1.0)
    p.add_run().add_picture(str(ROOT / "images/readme/dayglyph-logo.png"), width=Inches(0.72))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=2, line=1.08)
    run = p.add_run("DayGlyph 日迹：长期情绪画像驱动的自我观察应用")
    set_run_font(run, 19, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=8, line=1.1)
    run = p.add_run("情绪计算 / 生成式视觉叙事 / 本地优先隐私 / 低压力微行动 / 独立开发")
    set_run_font(run, 10, bold=True, color=BLUE)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    set_para(p, before=11, after=5, line=1.1)
    run = p.add_run(text)
    set_run_font(run, 14.5, bold=True, color=BLUE)


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para(p, after=5, line=1.2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    run = p.add_run(text)
    set_run_font(run, 10.3, color=DARK)


def add_caption(doc, text, after=7):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    set_para(p, before=0, after=after, line=1.0)
    run = p.add_run(text)
    set_run_font(run, 8.4, color=MUTED)


def add_figure(doc, rel_path, width_in, caption, before=2, after=7):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    set_para(p, before=before, after=1, line=1.0)
    p.add_run().add_picture(str(ROOT / rel_path), width=Inches(width_in))
    add_caption(doc, caption, after=after)


def add_cover_gallery(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=0, after=3, line=1.0)
    run = p.add_run("作品视觉总览")
    set_run_font(run, 11, bold=True, color=BLUE)

    add_figure(doc, "images/readme/feature-today.jpg", 6.45, "图 1  今日页：一句话生成情绪鸡尾酒、日星球与微行动", before=0, after=5)
    add_figure(doc, "images/readme/feature-universe.jpg", 6.45, "图 2  宇宙页：长期记录形成月星球、日期光点与趋势回顾", before=0, after=5)
    add_figure(doc, "images/readme/feature-echo.jpg", 3.0, "图 3  回声页：微行动后的真实反馈与个人发现", before=0, after=7)


def add_competitor_strip(doc, title, image_paths, width):
    p = doc.add_paragraph()
    set_para(p, before=4, after=2, line=1.0)
    run = p.add_run(title)
    set_run_font(run, 9.2, bold=True, color=DARK)

    table = doc.add_table(rows=1, cols=len(image_paths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_borders(table, color="EEF1F6", size="3")
    cell_margins(table, top=55, bottom=55, start=45, end=45)
    for cell, rel_path in zip(table.rows[0].cells, image_paths):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(paragraph, after=0, line=1.0)
        paragraph.add_run().add_picture(str(ROOT / rel_path), width=Inches(width))


def add_competitor_images(doc):
    add_competitor_strip(
        doc,
        "Daylio：轻量心情打卡与统计界面",
        [
            "竞品分析 image/dailo 1.png",
            "竞品分析 image/dailo2.png",
            "竞品分析 image/dailo3.png",
            "竞品分析 image/dailo4.png",
            "竞品分析 image/dailo5.png",
            "竞品分析 image/dailo6.png",
        ],
        0.69,
    )
    add_competitor_strip(
        doc,
        "Reflectly：AI 日记与情绪回顾界面",
        [
            "竞品分析 image/reflectly1.png",
            "竞品分析 image/reflectly 2.png",
        ],
        1.1,
    )
    add_competitor_strip(
        doc,
        "Stoic：结构化自我照护与练习界面",
        [
            "竞品分析 image/stocic1.png",
            "竞品分析 image/stoic 2.png",
            "竞品分析 image/stoic3.png",
            "竞品分析 image/stoic4.png",
        ],
        1.35,
    )
    add_caption(doc, "图 5  竞品分析界面样本：Daylio、Reflectly 与 Stoic", after=8)


def add_reference(doc, number, text, links):
    p = doc.add_paragraph()
    set_para(p, before=0, after=2, line=1.08)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.28)
    fmt.first_line_indent = Inches(-0.28)
    run = p.add_run(f"[{number}] ")
    set_run_font(run, 8.4, bold=True, color=DARK)
    run = p.add_run(text)
    set_run_font(run, 8.4, color=DARK)
    for label, url in links:
        p.add_run(" ")
        add_hyperlink(p, label, url)


def build():
    doc = Document()
    configure(doc)
    add_title(doc)
    add_cover_gallery(doc)

    add_body(
        doc,
        "DayGlyph 面向“未必需要临床服务、但需要看见和整理自己”的日常情绪场景。用户只需写下一句话，系统将其转化为多标签情绪结构、情绪鸡尾酒、日星球、微行动与长期情绪宇宙，让情绪从“难以命名的感觉”变成“可收藏、可回看、可轻量行动”的个人数据资产。",
    )

    add_heading(doc, "1、问题背景与用户分析（约200字）")
    add_figure(doc, "images/generated/problem-background-ai.png", 6.15, "图 4  问题背景信息图：年轻人日常情绪压力与低门槛自我观察需求", before=1, after=5)
    add_body(
        doc,
        "高校学生与年轻职场人常处在学业、实习、关系、家庭期待和自我评价交叠的压力场中；他们未必需要医疗化干预，却经常需要低门槛、低压力、可信任的自我观察入口。WHO 资料显示，全球约每 7 人中就有 1 人经历心理障碍[1]；表达性写作研究也表明，把模糊体验组织成语言，有助于个体整理压力经验[4]。现有日记工具要求长篇书写，情绪量表又容易把感受压扁成“好/坏/几分”。DayGlyph 的洞察是：当用户最累、最乱、最不想被评价时，最可持续的入口不是打卡和诊断，而是“一句话被温柔理解，并在之后的时间里被自己看见”。",
    )

    add_heading(doc, "2、相关竞品分析（约200字）")
    add_competitor_images(doc)
    add_body(
        doc,
        "Daylio 用心情图标和活动统计降低记录成本，但核心仍是手动选择与相关性图表，难以表达复杂、混合、矛盾的中文情绪[7]。Reflectly、Stoic 等 AI 日记或自我照护产品强调引导问题、习惯养成和正念练习，但容易进入“完成任务”的框架[8][9]；Apple 健康的 State of Mind 具备系统级隐私与情绪记录入口，但更像标准化日志[10]。DayGlyph 以自然语言为入口，用 VAD 连续情绪空间和 12 个情绪锚点承接复杂感受[2][3]，再用鸡尾酒与日星球建立记忆点；记录、回声和趋势默认本地保存，既保留 AI 表达力，也避免私密日记被平台化[5]。",
    )

    add_heading(doc, "3、可行性分析（约300字）")
    add_figure(doc, "images/generated/feasibility-architecture-ai.png", 6.1, "图 6  AI 生成架构图：DayGlyph 技术可行性闭环", before=1, after=5)
    add_body(
        doc,
        "DayGlyph 已具备端到端 iOS 原型。客户端采用 SwiftUI + SwiftData 构建今日、宇宙、回声、我的四 Tab；DayGenerationOrchestrator 用状态机编排“文本理解—Schema 校验—文本先展示—双图并行生成—本地保存—失败单项重试”，生成体验可恢复、可降级。AI 层使用豆包 Seed 2.0 Lite 输出结构化情绪 JSON，Seedream 生成情绪鸡尾酒与日星球；模型输出不直接控制 UI，而是经过 GenerationSchemaValidator、安全预检和受控提示词模板，降低幻觉、越界建议和视觉漂移风险。情绪模型采用 VAD 连续维度与 12 个 EmotionAnchor 的双重表示，兼容复杂情绪、统计聚合和历史稳定呈现。SwiftData 保存日记、分析、行动、回声和生成版本，图片落本地 Application Support；同一文本和日期通过稳定 seed 保持一致视觉。项目已有 Swift Testing、UI 测试、可访问降级、减少动态、VoiceOver 描述和高风险安全支持路径，具备比赛 Demo 到正式产品的工程基础。",
    )

    add_heading(doc, "4、App创新点（约300字）")
    add_figure(doc, "images/generated/innovation-loop-ai.png", 6.1, "图 7  AI 生成创新图：DayGlyph 创新体验闭环", before=1, after=5)
    add_body(
        doc,
        "DayGlyph 的创新不是“给日记加 AI”，而是一条可落地的情绪自我观察链路：一句话记录 → 情绪计算 → 生成式视觉印记 → 低压力微行动 → 行动回声 → 长程情绪宇宙。第一，把情绪从线性分数改写为可感知对象：同一次分析生成鸡尾酒与日星球，二者共享色板、材质、密度和光线方向，形成当天的“双视觉签名”。第二，支持 1～8 种复杂情绪的多标签理解，并投影到 VAD 连续空间和 12 个锚点，不强迫用户在“开心/难过/焦虑”间单选。第三，微行动是“可跳过、可失败、无惩罚”的三档建议，行动后只记录用户自己的回声，不承诺疗效、不推断因果。第四，情绪宇宙把每日记录聚合为月星球、日期光点、趋势图和文字替代，使长期数据既有审美吸引力，也有可访问与可解释能力。第五，本地优先、匿名共情改写校验、高风险短路和分享卡隐私控制，使敏感情绪数据从一开始就被当作用户资产。",
    )

    add_heading(doc, "5、应用前景（约200字）")
    add_figure(doc, "images/generated/prospect-roadmap-ai.png", 6.1, "图 8  AI 生成前景图：DayGlyph 应用前景路径", before=1, after=5)
    add_body(
        doc,
        "情绪记录与数字自我照护正在从“小众工具”变成高频需求。第三方研究预计，全球心理健康 App 市场将在 2030 年达到百亿美元级规模[6]；市场越大，用户越需要能守住隐私、不过度医疗化、不过度游戏化的产品。DayGlyph 的近期场景是高校学生、独立创作者、年轻职场人和咨询前后的自我观察记录，可作为个人情绪档案、咨询沟通前的自述材料、校园心理健康活动的低压力辅助工具。中期可扩展 iPad、Apple Watch、Shortcuts、本地导出、匿名分享卡和年度报告。长期看，DayGlyph 可形成“中文情绪表达词库 + 生成式视觉资产 + 自我观察方法论”的差异化壁垒，在不替代专业帮助的前提下，补足日常情绪觉察和长期回顾空白。",
    )

    add_heading(doc, "参考文献")
    refs = [
        ("World Health Organization. Mental disorders[EB/OL].", [("https://www.who.int/news-room/fact-sheets/detail/mental-disorders", "https://www.who.int/news-room/fact-sheets/detail/mental-disorders")]),
        ("Russell J A. A circumplex model of affect[J]. Journal of Personality and Social Psychology, 1980, 39(6): 1161-1178.", [("DOI:10.1037/h0077714", "https://doi.org/10.1037/h0077714")]),
        ("Bradley M M, Lang P J. Measuring emotion: The Self-Assessment Manikin and the semantic differential[J]. Journal of Behavior Therapy and Experimental Psychiatry, 1994, 25(1): 49-59.", [("PubMed", "https://pubmed.ncbi.nlm.nih.gov/7962581/")]),
        ("Smyth J M. Written emotional expression: Effect sizes, outcome types, and moderating variables[J]. Journal of Consulting and Clinical Psychology, 1998, 66(1): 174-184.", [("DOI:10.1037/0022-006X.66.1.174", "https://doi.org/10.1037/0022-006X.66.1.174")]),
        ("Kern J, et al. Privacy and security in mental health apps: A review of risks and design implications[EB/OL].", [("PMC", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9643945/"), ("Brookings", "https://www.brookings.edu/articles/why-mental-health-apps-need-to-take-privacy-more-seriously/")]),
        ("Grand View Research. Mental Health Apps Market Size, Share & Trends Analysis Report[EB/OL].", [("Grand View Research", "https://www.grandviewresearch.com/press-release/global-mental-health-apps-market")]),
        ("Daylio. Mood tracker and micro diary official website[EB/OL].", [("Daylio", "https://daylio.net/")]),
        ("Reflectly. Journal & AI diary official app information[EB/OL].", [("Reflectly", "https://apps.apple.com/us/app/reflectly-journal-ai-diary/id1241229134")]),
        ("Stoic. Journal & mental health official app information[EB/OL].", [("Stoic", "https://apps.apple.com/us/app/stoic-journal-mental-health/id1312926037")]),
        ("Apple. Log your state of mind in Health on iPhone[EB/OL].", [("Apple Support", "https://support.apple.com/guide/iphone/log-your-state-of-mind-iph6a6decb13/ios")]),
    ]
    for i, (text, links) in enumerate(refs, start=1):
        add_reference(doc, i, text, links)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
